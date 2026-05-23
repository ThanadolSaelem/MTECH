"""
Generate คู่มือการใช้งาน_MPTECH.docx
- Clones source docx (MTECH→MPTECH throughout)
- Replaces ALL existing images (paras 29/45/59/73) with fresh Xvfb screenshots
- Appends บทที่ 6 (Notifications), 7 (Error guide), 8 (Reference)
"""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image as PILImage

SRC = "/root/.claude/uploads/c2ca23fd-b630-4e72-9013-1ebfca2f7f24/c888ec48-________________MTECH.docx"
OUT = "/home/user/MTECH/คู่มือการใช้งาน_MPTECH.docx"
SS  = "/home/user/MTECH/manual_screenshots"

def ss(name):
    new = f"{SS}/ss_{name}_new.png"
    old = f"{SS}/ss_{name}.png"
    return new if os.path.exists(new) else old

SS_SET   = ss("settings")
SS_DASH  = ss("dashboard")
SS_TASKS = ss("tasks")
SS_LOGS  = ss("logs")
SS_NOTIF = f"{SS}/ss_notifications.png"

FONT  = "TH Sarabun New"
IMG_W = Cm(15)  # 5,400,000 EMU

# ── Image utilities ────────────────────────────────────────────────────────────

def _img_extent(path: str, width_emu: int) -> tuple[int, int]:
    """Return (cx, cy) maintaining aspect ratio."""
    with PILImage.open(path) as im:
        w, h = im.size
    return width_emu, int(width_emu * h / w)

def _add_image_rel(doc: Document, img_path: str) -> str:
    """Add image to document part, return rId."""
    rId, _image = doc.part.get_or_add_image(img_path)
    return rId

def replace_image_in_para(doc: Document, para_idx: int, img_path: str) -> None:
    """Replace the first image in paragraph[para_idx] with img_path."""
    p = doc.paragraphs[para_idx]

    # Find blip
    blip = p._p.find('.//' + qn('a:blip'))
    if blip is None:
        print(f"  WARNING: no blip in para {para_idx}")
        return

    # Find extent element to update dimensions
    ext = p._p.find('.//' + qn('wp:extent'))

    # Add new image relationship
    rId = _add_image_rel(doc, img_path)

    # Update blip embed
    blip.set(qn('r:embed'), rId)

    # Update extent dimensions
    if ext is not None:
        cx, cy = _img_extent(img_path, int(IMG_W))
        ext.set('cx', str(cx))
        ext.set('cy', str(cy))
        # Also update the xfrm ext inside spPr
        xfrm_ext = p._p.find('.//' + qn('a:ext'))
        if xfrm_ext is not None:
            xfrm_ext.set('cx', str(cx))
            xfrm_ext.set('cy', str(cy))

    # Update docPr name to avoid conflicts
    docPr = p._p.find('.//' + qn('wp:docPr'))
    if docPr is not None:
        docPr.set('name', os.path.basename(img_path))
    cNvPr = p._p.find('.//' + qn('pic:cNvPr'))
    if cNvPr is not None:
        cNvPr.set('name', os.path.basename(img_path))

    print(f"  ✓ para {para_idx} image replaced → {os.path.basename(img_path)}")

# ── Font helpers ───────────────────────────────────────────────────────────────

def _run_font(run, size_pt, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = FONT
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

def add_para(doc, text, size=14, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _run_font(run, size, bold, color)

def add_image(doc, path, caption=""):
    if not os.path.exists(path):
        add_para(doc, f"[ภาพ: {os.path.basename(path)} — ไม่พบไฟล์]",
                 size=12, color=RGBColor(0xb9, 0x1c, 0x1c))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=IMG_W)
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs:
            _run_font(r, 11, color=RGBColor(0x47, 0x55, 0x69))

def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def add_callout(doc, icon, text, bg=(0xff, 0xf3, 0xc7)):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "TableNormal"
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    _shade_cell(cell, "%02x%02x%02x" % bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tblBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "888888")
        tblBorders.append(b)
    tcPr.append(tblBorders)
    p = cell.paragraphs[0]
    run = p.add_run(f"{icon}  {text}")
    _run_font(run, 13)
    doc.add_paragraph()

def add_numbered_steps(doc, steps):
    tbl = doc.add_table(rows=len(steps), cols=2)
    tbl.style = "TableNormal"
    for i, step in enumerate(steps):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        c0.width = Cm(1.4)
        c1.width = Cm(17)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(str(i + 1))
        _run_font(r0, 13, bold=True)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(step)
        _run_font(r1, 13)
    doc.add_paragraph()

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "TableNormal"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        _shade_cell(cell, "dbeafe")
        r = cell.paragraphs[0].add_run(h)
        _run_font(r, 13, bold=True)
        if col_widths:
            cell.width = Cm(col_widths[j])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            _run_font(r, 13)
            if col_widths:
                cell.width = Cm(col_widths[j])
    doc.add_paragraph()

# ── Clone source + rename ──────────────────────────────────────────────────────

def clone_and_rename(src_path):
    doc = Document(src_path)
    for p in doc.paragraphs:
        for r in p.runs:
            if "MTECH" in r.text:
                r.text = r.text.replace("MTECH", "MPTECH")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if "MTECH" in r.text:
                            r.text = r.text.replace("MTECH", "MPTECH")
    doc.core_properties.title = "คู่มือการใช้งาน MPTECH ระบบบัญชี"
    return doc

# ── New chapters ──────────────────────────────────────────────────────────────

def build_new_chapters(doc):
    doc.add_page_break()

    # ── บทที่ 6 หน้าแจ้งเตือน ────────────────────────────────────────────────
    add_heading(doc, "บทที่ 6   หน้าแจ้งเตือน (Notifications)", level=1)
    add_para(doc,
        "หน้าแจ้งเตือนคือศูนย์กลางที่รวบรวมสิ่งที่ต้องทำและสถานะล่าสุดของระบบ "
        "— เปิดหน้านี้ก่อนทุกครั้งเมื่อเริ่มทำงาน")
    doc.add_paragraph()
    add_image(doc, SS_NOTIF, "หน้าแจ้งเตือน — ภาพรวม 4 ส่วน")

    add_heading(doc, "6.1  สัญลักษณ์ Badge บนเมนู", level=2)
    add_para(doc,
        "เมื่อมีรายการที่ต้องดูแล ตัวเลขสีแดงจะปรากฏบนปุ่ม \"แจ้งเตือน\" ในแถบ Sidebar "
        "— ระบบอัปเดต badge ทุก 90 วินาทีโดยอัตโนมัติ ไม่ต้องกด Refresh เอง")
    doc.add_paragraph()

    add_heading(doc, "6.2  ส่วนประกอบในหน้าแจ้งเตือน", level=2)
    add_table(doc,
        ["ส่วน", "สีหัวข้อ", "ความหมาย", "สิ่งที่ต้องทำ"],
        [
            ["Error ที่ต้องแก้", "แดง", "รายการที่ระบบสร้างเอกสารไม่สำเร็จ",
             "อ่านรายละเอียด แก้ไขตามคู่มือบทที่ 7"],
            ["งานที่ต้องลงมือ", "เหลือง/ส้ม",
             "รายการที่รอการดำเนินการจากเจ้าหน้าที่",
             "กลับไปหน้า Tasks แล้ว Run ตาม label ที่แสดง"],
            ["งานค้าง — ระบบทำต่อให้เองอัตโนมัติ", "น้ำเงิน",
             "Queue ที่รอผลจาก PEAK — ระบบ Poll ต่อเอง",
             "ไม่ต้องทำอะไร หรือกด Poll Queue ถ้าต้องการเร่ง"],
            ["สรุปกิจกรรมล่าสุด", "เทา",
             "ตารางสรุปผล Part ล่าสุดที่รันไป",
             "ตรวจตัวเลข Error — ถ้า > 0 ให้ดูหน้า Logs"],
        ],
        col_widths=[5.0, 2.8, 5.0, 5.5]
    )

    add_heading(doc, "6.3  วิธีใช้งาน", level=2)
    add_numbered_steps(doc, [
        "คลิก \"แจ้งเตือน\" ในแถบซ้าย (หรือดู badge สีแดงก่อน)",
        "อ่าน \"Error ที่ต้องแก้\" ก่อน — ถ้ามีรายการ ให้ทำตามบทที่ 7",
        "ตรวจ \"งานที่ต้องลงมือ\" — มักบอกว่าให้รัน Part ไหนต่อ",
        "ส่วน \"งานค้าง\" (สีน้ำเงิน) ไม่ต้องทำอะไร ระบบจัดการเอง",
        "กด Refresh เพื่อดึงข้อมูลล่าสุดทันที",
    ])
    add_callout(doc, "💡",
        "ถ้าหน้าแจ้งเตือนแสดง \"ไม่มีรายการ\" ทุกส่วน = ระบบทำงานปกติ ไม่มีอะไรต้องแก้ไข",
        bg=(0xdc, 0xfc, 0xe7))

    # ── บทที่ 7 คู่มือแก้ Error ─────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "บทที่ 7   คู่มือแก้ไขข้อผิดพลาด", level=1)
    add_para(doc,
        "เมื่อพบข้อความ Error ให้ค้นหาในตารางด้านล่าง แล้วทำตาม \"วิธีแก้ไข\" "
        "— ปัญหาส่วนใหญ่แก้ได้เองโดยไม่ต้องติดต่อทีม MPTECH")
    doc.add_paragraph()

    add_heading(doc, "7.1  ข้อผิดพลาดเรื่องการเชื่อมต่อ", level=2)
    add_table(doc,
        ["ข้อความ Error", "สาเหตุ", "วิธีแก้ไข"],
        [
            ["ยังไม่ได้ตั้งค่า GAS URL ในหน้า Settings",
             "ยังไม่กรอก URL หรือ URL ว่างเปล่า",
             "Settings → กรอก GAS Web App URL → Save Local → Test Connection"],
            ["ยังไม่ได้ตั้งค่า API Key ในหน้า Settings",
             "ยังไม่กรอก API Key",
             "Settings → กรอก API Key → Save Local"],
            ["เชื่อมต่ออินเทอร์เน็ตไม่ได้ — ตรวจสอบ WiFi",
             "อินเทอร์เน็ตขัดข้องหรือ WiFi หลุด",
             "1. ตรวจสอบ WiFi / สายแลน\n2. ทดสอบเปิดเว็บ browser\n3. กลับมากด Test Connection"],
            ["● Error หรือ ● Offline (มุมล่างซ้าย)",
             "เชื่อมต่อ GAS ไม่ได้ หรือ URL ผิด",
             "Settings → ตรวจ URL → กด Test Connection"],
        ],
        col_widths=[5.5, 4.5, 8.5]
    )

    add_heading(doc, "7.2  ข้อผิดพลาดเรื่องเวลา (Timeout)", level=2)
    add_table(doc,
        ["ข้อความ Error", "สาเหตุ", "วิธีแก้ไข"],
        [
            ["หมดเวลา (300s) — ลองใหม่อีกครั้ง",
             "งานเยอะมาก GAS ใช้เวลานาน แต่ระบบบันทึก checkpoint แล้ว",
             "กด \"Run อีกครั้ง\" ที่แถบเหลืองด้านล่าง Output\n"
             "ระบบจะเริ่มต่อจากจุดที่ค้างไว้ (ไม่ซ้ำรายการที่ทำแล้ว)\n"
             "รันซ้ำจนกว่า Output แสดง Error: 0"],
        ],
        col_widths=[5.5, 4.5, 8.5]
    )
    add_callout(doc, "📌",
        "ปุ่ม \"Run อีกครั้ง\" จะปรากฏอัตโนมัติหลังหมดเวลา — ไม่ต้องเลือก task ใหม่ กดปุ่มนี้ได้เลย",
        bg=(0xef, 0xf6, 0xff))

    add_heading(doc, "7.3  ข้อผิดพลาดจาก PEAK API", level=2)
    add_table(doc,
        ["ข้อความ Error", "สาเหตุ", "วิธีแก้ไข"],
        [
            ["PEAK API error: ยังไม่ได้ตั้งค่า CONNECT_ID / USER_TOKEN",
             "ยังไม่ Push Credentials ไปยัง GAS",
             "Settings → กรอก CONNECT_ID, USER_TOKEN, SPREADSHEET_ID → Push to GAS"],
            ["PEAK API 401 / Unauthorized",
             "USER_TOKEN หมดอายุ (Token มีอายุ 24 ชั่วโมง)",
             "ขอ Token ใหม่จาก PEAK → Settings → ใส่ USER_TOKEN ใหม่ → Push to GAS"],
            ["PEAK API 400: Transaction Limit exceeded",
             "ใช้เกิน quota API ของเดือน",
             "1. รอ 10 นาทีแล้วลองใหม่\n2. ถ้ายังเกิด แจ้งทีม MPTECH — อาจต้องอัปเกรด Package"],
            ["PEAK API 429: Too many requests",
             "ส่ง request ถี่เกินไป (rate limit)",
             "ระบบ retry อัตโนมัติแล้ว — ถ้ายังค้าง รอ 30 วินาทีแล้ว Run อีกครั้ง"],
            ["เลขที่เอกสารซ้ำ / Duplicate document",
             "เอกสารนี้ถูกสร้างใน PEAK แล้ว ระบบ detect ได้",
             "ระบบบันทึก [DUP] ให้อัตโนมัติ — งานเสร็จแล้ว ไม่ต้องทำอะไร"],
            ["ไม่พบ Contact / Contact not found",
             "สัญญา (invCode) ยังไม่ sync ไปยัง PEAK Contacts",
             "รัน Part นั้นซ้ำ — ระบบ sync Contact อัตโนมัติก่อนสร้างเอกสาร"],
        ],
        col_widths=[5.5, 4.5, 8.5]
    )

    add_heading(doc, "7.4  ข้อผิดพลาดเรื่อง Google Sheets", level=2)
    add_table(doc,
        ["อาการ", "สาเหตุ", "วิธีแก้ไข"],
        [
            ["Dashboard แสดง 0 ทุกการ์ด หรือ \"not found\"",
             "ชื่อ Sheet ไม่ตรง หรือยังไม่ได้เลือกเดือน",
             "กรอกเดือนให้ถูกรูปแบบ เช่น 05.2026 → กด Refresh"],
            ["ตัวเลขใน Dashboard ไม่อัปเดตหลัง Run",
             "ยังไม่ได้กด Refresh หรือ Queue ยังรอผล",
             "กด Refresh บน Dashboard — ถ้ามี Queue > 0 รอ Poll ก่อน"],
            ["Output แสดง Error: N รายการหลัง Run",
             "บางแถวยังออกเอกสารไม่ได้",
             "กด \"Run อีกครั้ง\" — ระบบข้ามรายการที่ทำแล้ว ทำแค่ส่วนที่ค้าง"],
            ["Part 5 Match Statement ไม่พบข้อมูล",
             "ชื่อ Sheet Statement หรือ Sum ไม่ถูกต้อง",
             "ในช่อง เดือน กรอก: ชื่อ Statement Sheet, ชื่อ Sum Sheet\nเช่น SCB05.2026,Sum05.2026"],
        ],
        col_widths=[5.5, 4.5, 8.5]
    )

    add_heading(doc, "7.5  Queue ค้าง", level=2)
    add_table(doc,
        ["อาการ", "วิธีแก้ไข"],
        [
            ["Queue > 0 แต่ไม่ลดลงนานหลายชั่วโมง",
             "Tasks → เลือก \"Poll Queue ทันที\" → Run"],
            ["Dashboard Queue ยังแสดงตัวเลขหลัง Poll",
             "PEAK ยังประมวลผลไม่เสร็จ — รอ 5-10 นาทีแล้ว Poll ซ้ำ"],
            ["Poll แล้ว Output แสดง \"Transaction Limit exceeded\"",
             "PEAK API quota เต็ม — รอ 10 นาทีแล้ว Poll ซ้ำ ระบบไม่ลบ Queue ไว้ให้"],
        ],
        col_widths=[7, 11.5]
    )

    add_heading(doc, "7.6  ขั้นตอนฉุกเฉิน (Emergency Checklist)", level=2)
    add_numbered_steps(doc, [
        "เปิดหน้า \"แจ้งเตือน\" → อ่านส่วน \"Error ที่ต้องแก้\" ก่อน",
        "เปิดหน้า \"Logs\" → กด Refresh → ดูคอลัมน์ Status และ msg ของแถวล่าสุด",
        "ค้นหาข้อความ Error ในตาราง 7.1–7.5 ด้านบน แล้วทำตาม",
        "ถ้าข้อความ Error ไม่อยู่ในตาราง — จดข้อความเต็ม ส่ง screenshot ให้ทีม MPTECH",
        "ห้ามลบหรือแก้ข้อมูลใน Google Sheets โดยตรง — แจ้งทีม MPTECH ก่อน",
    ])
    add_callout(doc, "⚠️",
        "ถ้าพบข้อความ [PROCESSING] ค้างในคอลัมน์ PEAK_DOC นานกว่า 1 ชั่วโมง "
        "แจ้งทีม MPTECH — อย่าลบออกเอง เพราะระบบใช้ค่านี้ป้องกันการสร้างเอกสารซ้ำ",
        bg=(0xff, 0xf1, 0xf1))

    # ── บทที่ 8 ข้อมูลอ้างอิง ────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "บทที่ 8   ข้อมูลอ้างอิง", level=1)

    add_heading(doc, "8.1  สัญลักษณ์สถานะในโปรแกรม", level=2)
    add_table(doc,
        ["สัญลักษณ์", "ความหมาย", "สิ่งที่ต้องทำ"],
        [
            ["● Connected (สีเขียว)", "เชื่อมต่อ GAS สำเร็จ — พร้อมใช้งาน", "ไม่ต้องทำอะไร"],
            ["● Offline (สีแดง)", "ไม่มีอินเทอร์เน็ต", "ตรวจ WiFi → กด Test Connection"],
            ["● Error (สีเหลือง/ส้ม)", "เชื่อมต่อได้แต่ GAS ตอบผิดปกติ", "กด Test Connection เพื่อดูรายละเอียด"],
            ["[PROCESSING]", "กำลังสร้างเอกสาร (ล็อคป้องกันซ้ำ)", "รอให้งานเสร็จ — ห้ามลบ"],
            ["[DUP]", "เอกสารซ้ำ — มีอยู่ใน PEAK แล้ว", "ปกติ ไม่ต้องทำอะไร"],
            ["[IN-PEAK]", "เอกสารอยู่ใน PEAK แต่ระบบไม่ทราบเลขที่", "แจ้งทีม MPTECH เพื่อกู้เลขที่เอกสาร"],
            ["TAX2026050001 / INV2026050001", "เลขที่เอกสารจาก PEAK — งานเสร็จ", "ตรวจสอบใน PEAK Account ได้เลย"],
        ],
        col_widths=[4.5, 7, 7]
    )

    add_heading(doc, "8.2  รายการ Tasks ทั้งหมด", level=2)
    add_table(doc,
        ["Task", "ใช้เมื่อไหร่", "ต้องใส่เดือน", "เวลาโดยประมาณ"],
        [
            ["Part 1 — ออกใบกำกับภาษี", "ต้นเดือน หลังรวบรวม Receipt ครบ", "✅", "3–10 นาที"],
            ["Part 1 — ค่าบริการเพิ่มเติม", "มีค่าบริการพิเศษนอกเหนือปกติ", "✅", "1–3 นาที"],
            ["Part 2 — ออกใบแจ้งหนี้ bulk", "มีสัญญาใหม่ที่ต้องออกใบแจ้งหนี้", "✅", "2–5 นาที"],
            ["Part 3 — ออกใบเสร็จค่าปรับ", "มีรายการค่าปรับ/ค่าชดเชย", "✅", "1–3 นาที"],
            ["Part 4 — ออกใบลดหนี้ (คืนเครื่อง)", "มีการคืนเครื่อง/ยกเลิกสัญญา", "❌", "1–2 นาที"],
            ["Part 5 — Match Statement", "ทุกเดือน หลังได้ Statement ธนาคาร", "✅ (2 ชีต)", "2–5 นาที"],
            ["Poll Queue ทันที", "หลัง Part 1/2/3 ถ้ามี Queue ค้าง", "❌", "< 1 นาที"],
            ["ทดสอบ PEAK Connection", "เมื่อสงสัยว่า PEAK API มีปัญหา", "❌", "< 30 วินาที"],
        ],
        col_widths=[5.0, 5.5, 3, 4]
    )

    add_heading(doc, "8.3  ข้อมูลการเชื่อมต่อ MPTECH", level=2)
    add_para(doc,
        "ข้อมูลด้านล่างนี้ใช้กรอกในหน้า Settings → Save Local "
        "เพื่อให้โปรแกรมเชื่อมต่อกับ GAS ได้", size=13)
    doc.add_paragraph()
    add_table(doc,
        ["รายการ", "ค่าที่ต้องกรอก"],
        [
            ["GAS Web App URL",
             "https://script.google.com/macros/s/AKfycbwnBO4IdcSyIH3RBxCSgLjsVuKdsS_"
             "Co3lCgZaZ2yCi8_XFo-PrYsZDC90tWW5dZNtm/exec"],
            ["API Key", "finfin-secret-2026"],
        ],
        col_widths=[4.0, 14.5]
    )
    add_callout(doc, "📌",
        "กรอกข้อมูลทั้ง 2 รายการในหน้า Settings → กด Save Local → กด Test Connection\n"
        "ถ้า ● Connected (สีเขียว) แสดงที่มุมล่างซ้าย = เชื่อมต่อสำเร็จ พร้อมใช้งาน",
        bg=(0xdc, 0xfc, 0xe7))

    add_heading(doc, "8.4  ผู้ดูแลระบบ", level=2)
    add_callout(doc, "📌",
        "ทีม MPTECH — ติดต่อผ่าน Line หรือโทรตรงเมื่อพบปัญหาที่แก้ตามคู่มือแล้วยังไม่หาย\n"
        "ข้อมูลที่ควรเตรียม: Screenshot + ข้อความ Error เต็มๆ + ชื่อ Task ที่รัน + วันเวลา",
        bg=(0xef, 0xf6, 0xff))


# ── Main ──────────────────────────────────────────────────────────────────────

print("1. Cloning source docx + renaming MTECH→MPTECH...")
doc = clone_and_rename(SRC)

print("2. Replacing existing screenshots...")
# para 29 = Settings (บทที่ 3)
# para 45 = Dashboard (บทที่ 4)
# para 59 = Tasks (บทที่ 5)
# para 73 = Logs (บทที่ 6 เก่า)
for para_idx, img_path, label in [
    (29, SS_SET,   "Settings"),
    (45, SS_DASH,  "Dashboard"),
    (59, SS_TASKS, "Tasks"),
    (73, SS_LOGS,  "Logs"),
]:
    if os.path.exists(img_path):
        replace_image_in_para(doc, para_idx, img_path)
    else:
        print(f"  SKIP {label}: file not found → {img_path}")

print("3. Appending new chapters...")
build_new_chapters(doc)

print(f"4. Saving → {OUT}")
doc.save(OUT)
sz = os.path.getsize(OUT)
print(f"Done! {sz:,} bytes")
